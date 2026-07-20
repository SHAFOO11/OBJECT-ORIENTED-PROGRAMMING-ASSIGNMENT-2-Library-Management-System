"""Generates BIT1123-Assignment2-Report.docx - a Word version of the same report
content as build_report.py, so the group can edit it directly (fill in member
details, tweak wording, etc.) without needing to touch the PDF pipeline.
"""

import re

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/shariiif77icloud.com/Desktop/BIT1123-Library-Management-System/report/BIT1123-Assignment2-Report.docx"
DIAGRAM = "/Users/shariiif77icloud.com/Desktop/BIT1123-Library-Management-System/report/uml_class_diagram.png"

DARK = RGBColor(0x1A, 0x1A, 0x2E)
SUBHEAD = RGBColor(0x4A, 0x4E, 0x69)
GREY = RGBColor(0x66, 0x66, 0x66)
CODE_BG = "F4F4F8"
CODE_BORDER = "DCDCE6"

ENTITIES = {
    "&ndash;": "–",
    "&rarr;": "→",
    "&ldquo;": "“",
    "&rdquo;": "”",
    "&lsquo;": "‘",
    "&rsquo;": "’",
}


def _decode_entities(text):
    for entity, char in ENTITIES.items():
        text = text.replace(entity, char)
    return text


def add_rich_runs(paragraph, html_text, base_size=10.5):
    """Parses a small subset of inline HTML (<font face='Courier'>, <i>, <b>)
    plus a handful of named entities, and writes styled runs into paragraph."""
    text = _decode_entities(html_text)
    tokens = re.split(r"(<font face='Courier'>|</font>|<i>|</i>|<b>|</b>)", text)
    courier = italic = bold = False
    for tok in tokens:
        if tok == "<font face='Courier'>":
            courier = True
            continue
        if tok == "</font>":
            courier = False
            continue
        if tok == "<i>":
            italic = True
            continue
        if tok == "</i>":
            italic = False
            continue
        if tok == "<b>":
            bold = True
            continue
        if tok == "</b>":
            bold = False
            continue
        if not tok:
            continue
        run = paragraph.add_run(tok)
        run.font.size = Pt(base_size)
        run.italic = italic
        run.bold = bold
        if courier:
            run.font.name = "Courier New"
            run.font.color.rgb = DARK
        else:
            run.font.name = "Helvetica"


def body_paragraph(doc, html_text, space_after=9):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.3
    add_rich_runs(p, html_text)
    return p


def section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = "Helvetica"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = DARK
    return p


def sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Helvetica"
    run.font.size = Pt(12.5)
    run.font.bold = True
    run.font.color.rgb = SUBHEAD
    return p


def _shade_cell(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_border(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), hex_color)
        borders.append(el)
    tcPr.append(borders)


def code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Cm(17)
    _shade_cell(cell, CODE_BG)
    _set_cell_border(cell, CODE_BORDER)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    lines = text.split("\n")
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = DARK
    for row in table.rows:
        row.cells[0].paragraphs[0].paragraph_format.space_before = Pt(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        add_rich_runs(p, item)


def add_page_break(doc):
    doc.add_page_break()


def set_footer_page_number(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    run.text = "Page "

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run2 = p.add_run()
    run2.font.size = Pt(8)
    run2.font.color.rgb = GREY
    run2._r.append(fld_begin)
    run2._r.append(instr)
    run2._r.append(fld_end)


doc = Document()

section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

set_footer_page_number(doc)

# ---------------- Cover Page ----------------
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BIT1123 / BISE2093 / DIT1113")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("OBJECT ORIENTED PROGRAMMING")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ASSIGNMENT 2 (GROUP) – 20%")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Library Management System")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = DARK

for _ in range(2):
    doc.add_paragraph()

member_data = [["Full Name", "Student ID", "Class Code", "Program", "NRIC / Passport No."]]
for _ in range(4):
    member_data.append(["", "", "", "", ""])

table = doc.add_table(rows=len(member_data), cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for col in table.columns:
    for cell in col.cells:
        cell.width = Cm(3.2)
for r, row_data in enumerate(member_data):
    for c, value in enumerate(row_data):
        cell = table.rows[r].cells[c]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(value)
        run.font.size = Pt(8.5)
        _set_cell_border(cell, "CCCCCC")
        if r == 0:
            _shade_cell(cell, "1A1A2E")
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("(Fill in the details of every group member above before submission.)")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = GREY

for _ in range(3):
    doc.add_paragraph()

body_paragraph(doc, "<b>GitHub Repository:</b> [ paste your repository URL here ]")
body_paragraph(doc, "<b>Lecturer:</b> Sir Nazmirul Izzad Bin Nassir")
body_paragraph(doc, "<b>Faculty:</b> Faculty of Information Technology, City University Malaysia, Cyberjaya Campus")
add_page_break(doc)

# ---------------- 1. Introduction ----------------
section_heading(doc, "1. Introduction")
body_paragraph(doc,
    "For this assignment we were asked to design and build a Java console application that puts "
    "Object-Oriented Programming into practice rather than just describing it on paper. The brief "
    "specifically calls for classes and objects, encapsulation, inheritance, polymorphism, and "
    "abstraction to all show up somewhere in a real, working program, and for the report to explain "
    "how and why each one was used. We chose to build a Library Management System, and this report "
    "walks through the reasoning behind that choice, how the classes fit together, and how each of "
    "the five OOP requirements is actually expressed in the code, not just claimed about it.")
body_paragraph(doc,
    "We didn't want to just tick boxes on a rubric. Before writing any code we talked through what a "
    "library actually needs to keep track of, sketched out which real-world things should become "
    "classes, and only then started implementing. The result is a program with nine classes that a "
    "librarian can run from the terminal to register members and staff, add books and DVDs to the "
    "catalog, lend items out, take them back and work out any late fee, search the catalog, and save "
    "everything to a file so nothing is lost when the program is closed. Every part of that workflow "
    "is described in the sections that follow, together with the class diagram and the specific lines "
    "of code that demonstrate each OOP concept.")

# ---------------- 2. Problem Description ----------------
section_heading(doc, "2. Problem Description")
body_paragraph(doc,
    "Picture a small college library with one or two staff on the front desk. A student walks up with "
    "a book, the staff member writes the student's name and the book's title into a notebook or an "
    "Excel sheet, and moves on to the next person. This works fine until the notebook fills up, "
    "someone forgets to write down the date, or two different staff members note the same book as "
    "borrowed by two different people because nobody had a single, shared source of truth. Late fees "
    "are worse: a DVD that comes back five days late should probably cost more per day than a "
    "paperback five days late, but if the fee is worked out by hand it depends entirely on whoever is "
    "on duty remembering the right rate.")
body_paragraph(doc,
    "That's really the core problem we set out to solve: keep track of three things that are always "
    "moving – who is registered at the library, what items exist and whether they're currently "
    "available, and which member currently holds which item – in a way that doesn't rely on "
    "someone's memory or a shared spreadsheet. Object-oriented design turned out to be a natural fit "
    "for this, because each of those three concerns maps cleanly onto its own class with its own data "
    "and its own rules, instead of one giant function juggling arrays of loosely related values.")
body_paragraph(doc, "Concretely, the console application we built lets a librarian:")
bullet_list(doc, [
    "Register new members and new librarians, each keeping their own contact details and "
    "role-specific information",
    "Add new books and DVDs to the catalog, each with the details relevant to that type of item",
    "Print out a formatted list of every member, librarian, catalog item, or loan currently on record",
    "Borrow and return items, with the late fee worked out automatically and charged at a different "
    "daily rate depending on whether the item is a book or a DVD",
    "Search the catalog by typing part of a title",
    "Save the current state of the library to a text file, and load it back in on the next run, so "
    "closing the program doesn't wipe out the day's work",
])

# ---------------- 3. Class Diagram ----------------
add_page_break(doc)
section_heading(doc, "3. Class Diagram (UML)")
body_paragraph(doc,
    "Figure 1 shows all nine classes in the system with their attributes and methods, using standard "
    "UML visibility notation (<font face='Courier'>-</font> for private, <font face='Courier'>+</font> "
    "for public). The arrows carry meaning too: a hollow triangle points from a subclass up to its "
    "superclass (inheritance), an open diamond sits on the side of the class that owns a collection of "
    "the other (aggregation), a plain arrow shows one class referring to another (association), and the "
    "dashed arrow from <font face='Courier'>Main</font> shows a dependency rather than a stored "
    "reference. We laid the diagram out by hand so the two inheritance branches sit clearly apart on "
    "the left and right, with <font face='Courier'>Library</font> and <font face='Courier'>Loan</font> "
    "bridging them down the middle – it took a couple of attempts to get the lines routed "
    "without them crossing through unrelated boxes, but the payoff is that the relationships are "
    "actually easy to trace by eye.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture(DIAGRAM, width=Cm(16.5))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Figure 1: UML Class Diagram of the Library Management System")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = GREY

# ---------------- 4. UML Design Explanation ----------------
section_heading(doc, "4. Explanation of UML Design")
body_paragraph(doc,
    "The design centres on two separate inheritance hierarchies, connected through an association "
    "class and managed by a single controller class. None of this was the first idea we had – "
    "we went through a couple of rounds of sketching before settling on the structure shown in Figure "
    "1, and it's worth explaining what we tried and rejected along the way, since that reasoning is "
    "really the point of a design justification section.")

sub_heading(doc, "Person → Member / Librarian")
body_paragraph(doc,
    "Our first instinct was actually a single <font face='Courier'>Person</font> class with a "
    "<font face='Courier'>role</font> field that could be set to “member” or "
    "“librarian”, and a pile of if-statements checking which one it was. We dropped that "
    "almost immediately because it's exactly the kind of design inheritance exists to replace: a "
    "librarian doesn't have a fine balance or a borrowed-item count, and a member doesn't have a "
    "department, so cramming both into one class means every object carries fields that don't apply "
    "to it. Instead, <font face='Courier'>Person</font> holds only what every person in the system "
    "genuinely shares – a name, an ID, and a contact number – and is declared abstract "
    "because “a person” on its own is never actually registered at the library; only a "
    "member or a librarian is. <font face='Courier'>Member</font> then adds membership ID, borrowed "
    "count, and fine balance, while <font face='Courier'>Librarian</font> adds staff ID and "
    "department. Each subclass calls <font face='Courier'>super(...)</font> to set up the shared "
    "fields and then only worries about what makes it different.")

sub_heading(doc, "LibraryItem → Book / DVD")
body_paragraph(doc,
    "The same reasoning applies on the catalog side, but with an extra wrinkle: books and DVDs aren't "
    "just different in what data they store (author and ISBN versus director and running time), they "
    "also need to be charged late fees at different daily rates. That's precisely the situation an "
    "abstract method is built for. <font face='Courier'>LibraryItem</font> declares "
    "<font face='Courier'>calculateLateFee(daysLate)</font> and <font face='Courier'>getItemType()</font> "
    "but deliberately provides no body for either – it can't, because it has no idea what rate "
    "applies until you know whether the concrete object is a book or a DVD. <font face='Courier'>Book</font> "
    "charges RM0.50 per day late; <font face='Courier'>DVD</font> charges RM1.00 per day late. Adding "
    "a third item type later, say a <font face='Courier'>Magazine</font>, would mean writing one new "
    "class that extends <font face='Courier'>LibraryItem</font> and nothing else in the system would "
    "need to change.")

sub_heading(doc, "Loan as an association class")
body_paragraph(doc,
    "We also considered just giving <font face='Courier'>Member</font> a list of the items it "
    "currently has out, and skipping a separate loan class entirely. That falls apart quickly once "
    "you think about late fees and history: a fine has to be calculated from a specific borrow date, "
    "and a returned loan is still worth keeping a record of, even after the item goes back on the "
    "shelf. Bolting all of that onto <font face='Courier'>Member</font> would tie it tightly to "
    "<font face='Courier'>LibraryItem</font> and make the two classes depend on each other directly. "
    "So instead we introduced <font face='Courier'>Loan</font> as its own class that sits between "
    "them, holding a reference to the <font face='Courier'>Member</font> who borrowed something, a "
    "reference to the <font face='Courier'>LibraryItem</font> they borrowed, the date it was borrowed, "
    "and whether it's been returned. Neither <font face='Courier'>Member</font> nor "
    "<font face='Courier'>LibraryItem</font> needs to know the other exists – "
    "<font face='Courier'>Loan</font> is the only class that connects them, which is exactly what an "
    "association class is for.")

sub_heading(doc, "Library as the controller")
body_paragraph(doc,
    "Finally, something had to actually own the four growing lists of members, librarians, items, and "
    "loans, and provide the operations that act on them – registering people, borrowing and "
    "returning items, searching, saving and loading. We put all of that in "
    "<font face='Courier'>Library</font> rather than scattering it across "
    "<font face='Courier'>Main</font>, so that <font face='Courier'>Main</font> only has to worry "
    "about reading console input and printing results. That's why the diagram shows "
    "<font face='Courier'>Main</font> depending on <font face='Courier'>Library</font> with a dashed "
    "arrow rather than the two being tightly coupled: if we ever swapped the console menu for, say, a "
    "simple GUI, <font face='Courier'>Library</font> wouldn't need to change at all.")

# ---------------- 5. Explanation of OOP Concepts Used ----------------
add_page_break(doc)
section_heading(doc, "5. Explanation of OOP Concepts Used")
body_paragraph(doc,
    "This section goes through each of the five OOP requirements one at a time, pointing at the exact "
    "class or method where it shows up and, where it makes sense, at what we actually saw happen when "
    "we ran the program.")

sub_heading(doc, "5.1 Classes and Objects")
body_paragraph(doc,
    "There are nine classes in total: <font face='Courier'>Person</font>, "
    "<font face='Courier'>Member</font>, <font face='Courier'>Librarian</font>, "
    "<font face='Courier'>LibraryItem</font>, <font face='Courier'>Book</font>, "
    "<font face='Courier'>DVD</font>, <font face='Courier'>Loan</font>, "
    "<font face='Courier'>Library</font>, and <font face='Courier'>Main</font>, and each one models a "
    "single real-world idea rather than being a grab-bag of unrelated data. A class is just the "
    "blueprint, though – nothing actually exists until the menu creates an object from it. When "
    "a librarian picks “Add new book” from the menu and types in a title and author, "
    "<font face='Courier'>Main</font> creates one specific <font face='Courier'>Book</font> object "
    "holding those exact values, completely separate from every other book already in the catalog.")
body_paragraph(doc, "This is what that object creation looks like in Main.java:")
code_block(doc,
    "library.registerMember(new Member(name, id, contact, membershipId));\n"
    "library.addItem(new Book(itemId, title, author, isbn));")

sub_heading(doc, "5.2 Encapsulation")
body_paragraph(doc,
    "Every field in every class is private. Nothing outside the class can reach in and change an "
    "object's data directly – the only way in or out is through the public getter and setter "
    "methods the class chooses to expose, and those methods can enforce rules on the way. "
    "<font face='Courier'>Member.addFine()</font> is a small but telling example: it only accepts a "
    "positive amount, so there's no way for a bug elsewhere in the program to accidentally push a "
    "member's fine balance negative. If <font face='Courier'>fineBalance</font> were a public field, "
    "any piece of code anywhere could set it to whatever it wanted, including nonsense values, and we "
    "wouldn't be able to guarantee that check ever ran.")
code_block(doc,
    "public double getFineBalance() {\n"
    "    return fineBalance;\n"
    "}\n\n"
    "public void addFine(double amount) {\n"
    "    if (amount > 0) {\n"
    "        this.fineBalance += amount;\n"
    "    }\n"
    "}")

sub_heading(doc, "5.3 Inheritance")
body_paragraph(doc,
    "Both hierarchies discussed in Section 4 use Java's <font face='Courier'>extends</font> keyword: "
    "<font face='Courier'>Member</font> and <font face='Courier'>Librarian</font> extend "
    "<font face='Courier'>Person</font>, and <font face='Courier'>Book</font> and "
    "<font face='Courier'>DVD</font> extend <font face='Courier'>LibraryItem</font>. In every case the "
    "subclass constructor calls <font face='Courier'>super(...)</font> first to let the parent class "
    "set up the fields it owns, and only then initializes its own additional fields. It's worth "
    "pointing out that <font face='Courier'>Member</font> never re-declares "
    "<font face='Courier'>name</font>, <font face='Courier'>id</font>, or "
    "<font face='Courier'>contactNumber</font> – it inherits all three from "
    "<font face='Courier'>Person</font> and simply reuses <font face='Courier'>getName()</font> and "
    "the rest without writing a single extra line for them.")
code_block(doc,
    "public class Member extends Person {\n\n"
    "    private String membershipId;\n"
    "    private int borrowedCount;\n"
    "    private double fineBalance;\n\n"
    "    public Member(String name, String id, String contactNumber, String membershipId) {\n"
    "        super(name, id, contactNumber);\n"
    "        this.membershipId = membershipId;\n"
    "        this.borrowedCount = 0;\n"
    "        this.fineBalance = 0.0;\n"
    "    }\n"
    "}")

sub_heading(doc, "5.4 Polymorphism")
body_paragraph(doc,
    "The clearest example of runtime polymorphism in the project is how a late fee actually gets "
    "calculated. <font face='Courier'>Loan.calculateFine()</font> never checks whether it's holding a "
    "book or a DVD – it just calls <font face='Courier'>item.calculateLateFee(daysLate)</font> "
    "on whatever <font face='Courier'>LibraryItem</font> reference it has, and lets Java figure out at "
    "runtime which override actually runs. The same idea shows up again in "
    "<font face='Courier'>Library.displayAllItems()</font>, which loops over every item using the "
    "superclass type and calls <font face='Courier'>displaySummary()</font> on each one without caring "
    "what's underneath.")
code_block(doc,
    "public double calculateFine(int daysLate) {\n"
    "    return item.calculateLateFee(daysLate);\n"
    "}")
body_paragraph(doc,
    "We didn't just take this on faith – we tested it directly while trying out the program. "
    "Borrowing a book and returning it three days late charged RM1.50 (three days at RM0.50), while "
    "borrowing a DVD and returning it three days late through that exact same "
    "<font face='Courier'>returnItem()</font> code path charged RM3.00 (three days at RM1.00). Same "
    "method call, same loop, two different results, because the object underneath the "
    "<font face='Courier'>LibraryItem</font> reference was different each time. That's polymorphism "
    "actually doing something useful rather than just being a term we can define.")

sub_heading(doc, "5.5 Abstraction")
body_paragraph(doc,
    "<font face='Courier'>Person</font> and <font face='Courier'>LibraryItem</font> are both declared "
    "<font face='Courier'>abstract</font>, which means Java won't let anyone write "
    "<font face='Courier'>new Person(...)</font> or <font face='Courier'>new LibraryItem(...)</font> "
    "anywhere in the code – only their concrete subclasses can be instantiated. That's a "
    "deliberate restriction, not a limitation: on their own, “a person” or “a "
    "library item” are too vague to have a working "
    "<font face='Courier'>calculateLateFee()</font> or a working "
    "<font face='Courier'>displayDetails()</font>, so the abstract class only defines the shared shape "
    "and leaves the actual behaviour for each subclass to fill in.")
code_block(doc,
    "public abstract class LibraryItem {\n\n"
    "    private String itemId;\n"
    "    private String title;\n"
    "    private boolean available;\n\n"
    "    public abstract double calculateLateFee(int daysLate);\n\n"
    "    public abstract String getItemType();\n"
    "}")

# ---------------- 6. Sample Output ----------------
add_page_break(doc)
section_heading(doc, "6. Sample Output")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run(
    "[ Insert screenshots of your own terminal run here before submission. A good set to include: "
    "the main menu, registering a member, adding a book and a DVD, borrowing an item, returning a book "
    "late versus returning a DVD late side by side (so the different fine amounts are visible), and "
    "the save/load confirmation messages. ]")
run.italic = True
run.font.size = Pt(10.5)
body_paragraph(doc, "To compile and run the program from the project root:")
code_block(doc,
    "cd BIT1123-Library-Management-System\n"
    "javac -d out src/*.java\n"
    "java -cp out Main")

# ---------------- 7. Conclusion ----------------
section_heading(doc, "7. Conclusion")
body_paragraph(doc,
    "Working through this assignment made the four pillars of OOP feel a lot less like textbook "
    "vocabulary and more like decisions that actually change how easy a program is to extend. The "
    "clearest moment of that was realizing how little would need to change if we wanted to add a new "
    "kind of item to the catalog: one new class extending <font face='Courier'>LibraryItem</font>, "
    "and every part of the system that already loops over items – searching, displaying, "
    "borrowing, returning – would handle it automatically through polymorphism, without a single "
    "if-statement checking “is this a book or a DVD” anywhere in "
    "<font face='Courier'>Library</font> or <font face='Courier'>Main</font>.")
body_paragraph(doc,
    "The design wasn't perfect on the first attempt, and we don't think it needs to be presented as "
    "if it was – the single-class-with-a-role-field idea we tried early on for "
    "<font face='Courier'>Person</font>, and the version of <font face='Courier'>Library</font> that "
    "originally didn't persist active loans at all, both got reworked once we actually tested the "
    "program end to end and saw where they broke down. If we kept building on this, the next things "
    "we'd want to add are due dates that are checked automatically rather than typed in by the "
    "librarian, and proper authentication so that only a logged-in librarian account can register new "
    "members or add items. For a first pass at applying OOP to a real problem, though, we're happy "
    "with how cleanly the final structure maps onto the actual rules of running a library.")

doc.save(OUT)
print("DOCX written to", OUT)
